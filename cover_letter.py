import os
import re
from pathlib import Path
import anthropic
from docx import Document
from dotenv import load_dotenv

load_dotenv()

client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
COVER_LETTERS_DIR = Path(__file__).parent / "cover_letters"
BASE_RESUME_PATH = Path(__file__).parent / "resume.docx"


def _read_docx(path: Path) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _safe_filename(text: str) -> str:
    return re.sub(r"[^\w\-]", "_", text.strip()).lower()


def _write_docx(content: str, path: Path) -> None:
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)


def generate_cover_letter(job: dict) -> Path:
    """Generate a personalised cover letter grounded in the real resume."""
    base_text = _read_docx(BASE_RESUME_PATH)
    jd = job["description"]
    company = _safe_filename(job["company"])
    title = _safe_filename(job["title"])
    out_path = COVER_LETTERS_DIR / f"{company}_{title}_coverletter.docx"

    prompt = f"""You are a professional cover letter writer. Write a compelling, authentic cover letter
for the candidate applying to the role below.

STRICT RULES:
1. Base EVERY claim strictly on the candidate's real experience in the resume below.
2. NEVER fabricate skills, projects, or achievements.
3. Mirror the language and key terms from the job description naturally.
4. Keep the tone professional, confident, and human — not robotic.
5. Structure: opening hook (1 short paragraph), 2 body paragraphs (relevant experience + why this company),
   closing paragraph (call to action). No more than 400 words total.
6. Do NOT include placeholder text like "[Your Name]" — use the name found in the resume.
7. Address it to the Hiring Team at {job["company"]} if no specific contact is available.

CANDIDATE RESUME:
{base_text}

JOB TITLE: {job["title"]}
COMPANY: {job["company"]}
LOCATION: {job["location"]}

JOB DESCRIPTION:
{jd}

Write the cover letter below:"""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        messages=[{"role": "user", "content": prompt}],
    )

    letter_text = message.content[0].text
    _write_docx(letter_text, out_path)
    print(f"[cover_letter] Saved cover letter → {out_path.name}")
    return out_path
