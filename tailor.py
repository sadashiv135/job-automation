import os
import re
import shutil
from pathlib import Path
import anthropic
from docx import Document
from dotenv import load_dotenv

load_dotenv()

client           = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
RESUMES_DIR      = Path(__file__).parent / "resumes"
BASE_RESUME_PATH = Path(os.environ.get("RESUME_PATH", Path(__file__).parent / "resume.docx"))


def _safe_filename(text: str) -> str:
    return re.sub(r"[^\w\-]", "_", text.strip()).lower()


def _is_contact_para(text: str) -> bool:
    """Return True for name/contact header lines that must never be rewritten."""
    t = text.lower()
    return bool(
        re.search(r'\d{3}[-.\s]\d{3}[-.\s]\d{4}', text) or  # phone
        re.search(r'[\w.+-]+@[\w-]+\.\w+', text) or          # email
        'github' in t or
        'linkedin' in t
    )


def _parse_numbered_lines(response: str, expected: int) -> list[str]:
    """Parse 'N: text' lines from Claude's response into a list indexed from 0."""
    result: dict[int, str] = {}
    for line in response.strip().split('\n'):
        m = re.match(r'^(\d+):\s*(.*)', line.strip())
        if m:
            result[int(m.group(1))] = m.group(2)
    return [result.get(i, '') for i in range(1, expected + 1)]


def _write_docx(tailored_lines: list[str], body_indices: list[int],
                original_path: Path, out_path: Path) -> None:
    """
    Copy original_path to out_path, then overwrite only the body paragraphs
    (identified by body_indices) with tailored_lines, preserving run formatting.
    Header/contact paragraphs are left untouched.
    """
    shutil.copy2(original_path, out_path)
    doc = Document(out_path)

    for para_idx, new_text in zip(body_indices, tailored_lines):
        if not new_text:
            continue
        para = doc.paragraphs[para_idx]
        if not para.runs:
            continue
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''

    doc.save(out_path)


def tailor_resume(job: dict) -> Path:
    """Tailor base resume to the job description. Returns path to tailored .docx."""
    os.makedirs(RESUMES_DIR, exist_ok=True)

    orig_doc = Document(BASE_RESUME_PATH)

    # Split paragraphs into contact (keep verbatim) and body (send to Claude)
    nonempty = [(i, p.text) for i, p in enumerate(orig_doc.paragraphs) if p.text.strip()]
    body     = [(i, text) for i, text in nonempty if not _is_contact_para(text)]

    body_indices = [i for i, _ in body]
    numbered_in  = '\n'.join(f"{n}: {text}" for n, (_, text) in enumerate(body, 1))
    total_lines  = len(body)

    company  = _safe_filename(job["company"])
    title    = _safe_filename(job["title"])
    out_path = RESUMES_DIR / f"{company}_{title}_resume.docx"

    prompt = f"""You are a professional resume editor. Tailor the candidate's resume to better \
match the job description WITHOUT fabricating experience.

RULES (strictly follow):
1. Keep all experience stories, projects, and achievements exactly as-is.
2. You MAY adjust keyword phrasing to mirror the JD where the underlying skill is the same.
3. You MAY reword existing bullets to use JD terminology.
4. You MAY reorder bullets within a section to surface the most relevant items first.
5. NEVER add a skill, tool, or achievement not already in the resume.
6. NEVER remove experience sections.
7. Return EXACTLY {total_lines} numbered lines in the format "N: text". One line per number.
   Do NOT add, remove, merge, or split lines.

ONE PAGE CONSTRAINT:
The tailored resume MUST fit on one page. Achieve this through CONCISENESS, not by removing lines.
Return EXACTLY {total_lines} numbered lines — do not skip or omit any number.
- Aim for one tight sentence per bullet (20 words or fewer).
- NEVER reorder content across sections (do not swap project order, do not move section headers).
  You may only reorder bullets WITHIN the same job role.

BULLET POINT FORMAT — STAR METHOD:
Rewrite all experience bullet points using the STAR method:
  Situation/Task + Action + Result with metrics.
Format: "[Action verb] [what you built/did] using [tech stack] to [solve what problem], resulting in [measurable outcome]"
- Every bullet must have a quantifiable result where possible.
- If no metric exists use relative improvement language: "significantly", "substantially", "measurably improved".

RESUME ({total_lines} numbered lines):
{numbered_in}

JOB DESCRIPTION:
{job["description"]}

Return ONLY the {total_lines} numbered lines. No explanations, no blank lines between entries."""

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    tailored_lines = _parse_numbered_lines(message.content[0].text, total_lines)
    _write_docx(tailored_lines, body_indices, BASE_RESUME_PATH, out_path)
    print(f"[tailor] Saved tailored resume → {out_path.name}")
    return out_path
